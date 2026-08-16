"""
docx_builder.py
----------------
Tayyor savol/variant ro'yxatidan (Streamlit'da tahrirlangan yoki bot orqali
to'g'ridan-to'g'ri kelgan) yakuniy Word (.docx) faylni yasaydi.

Format aynan namunadagi (YHQ_test_savollari.docx) kabi:

    YHQ Test Savollari        <- sarlavha (Heading 1, markazda)

    [ASL RASM]                <- savolga tegishli original skrinshot/rasm

    1). Savol matni            <- qalin (bold)
    A) Variant
    B) Variant
    C) Variant
    *D) To'g'ri variant        <- qalin (bold), boshida * belgisi
    (bo'sh qator)

    2). Keyingi savol ...
"""

import io
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FKEY_TO_LETTER = ['A', 'B', 'C', 'D', 'E', 'F']

# Namunadagi shrift o'lchami (152400 EMU/12700 = 12pt) va rasm kengligi
BODY_FONT_SIZE = Pt(12)
IMAGE_WIDTH_INCHES = Inches(3.65)


def build_docx(questions, output_path, title="Test Savollari"):
    """questions: [{
           "question": str,
           "options": [str, ...],
           "correct_index": int|None,
           "image_bytes": bytes|None,   # <-- asl rasm (ixtiyoriy)
       }, ...]
    output_path: yakuniy .docx fayl saqlanadigan yo'l."""
    doc = Document()

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, q in enumerate(questions, start=1):
        question_text = q.get("question", "").strip()
        options = q.get("options", [])
        correct_index = q.get("correct_index")
        image_bytes = q.get("image_bytes")
        if image_bytes is None and q.get("image_b64"):
            try:
                image_bytes = base64.b64decode(q["image_b64"])
            except Exception:
                image_bytes = None

        if not question_text or not options:
            continue

        # --- Asl rasm (agar mavjud bo'lsa) - savol matnidan OLDIN, markazda ---
        if image_bytes:
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                img_paragraph.add_run().add_picture(
                    io.BytesIO(image_bytes), width=IMAGE_WIDTH_INCHES
                )
            except Exception:
                # Rasm buzuq/o'qib bo'lmasa - shunchaki rasmsiz davom etamiz
                pass

        # --- Savol matni (qalin) ---
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{i}). {question_text}")
        q_run.bold = True
        q_run.font.size = BODY_FONT_SIZE

        # --- Variantlar ---
        for j, option_text in enumerate(options):
            letter = FKEY_TO_LETTER[j] if j < len(FKEY_TO_LETTER) else str(j + 1)
            is_correct = correct_index == j
            prefix = f"*{letter})" if is_correct else f"{letter})"

            opt_para = doc.add_paragraph()
            opt_run = opt_para.add_run(f"{prefix} {option_text.strip()}")
            opt_run.font.size = BODY_FONT_SIZE
            if is_correct:
                opt_run.bold = True

        doc.add_paragraph()

    doc.save(output_path)
    return output_path
